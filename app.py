import streamlit as st
import anthropic
import json
import io
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Configuración ──────────────────────────────────────────────────────────────
st.set_page_config(page_title="Generador de Actas", layout="centered")
EMPRESAS = ["ESTILO GOURMET", "MUNDIPRODUCTOS", "PERSONA U OTRA"]
client = anthropic.Anthropic(api_key=st.secrets["ANTHROPIC_API_KEY"])

# ── Helpers DOCX ───────────────────────────────────────────────────────────────
def set_font(run, size=10, bold=False):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = None  # negro

def add_paragraph(doc, text="", bold=False, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=0):
    p = doc.add_paragraph()
    p.alignment = align
    fmt = p.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    fmt.line_spacing = Pt(12)
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold)
    return p

def add_bullet(doc, text, size=10):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = p.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = Pt(12)
    run = p.add_run(text)
    set_font(run, size=size)
    return p

def add_hrule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_section_title(doc, text):
    add_paragraph(doc, text, bold=True, size=10, space_before=6, space_after=0)
    add_hrule(doc)

def build_docx(acta, participantes, empresa, elaborado_por):
    doc = Document()

    # Tamaño carta y márgenes
    section = doc.sections[0]
    section.page_width  = Cm(21.59)
    section.page_height = Cm(27.94)
    section.top_margin    = Cm(3.0)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

    # Sin encabezado ni pie de página
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    for p in section.header.paragraphs:
        for r in p.runs:
            r.text = ""
    for p in section.footer.paragraphs:
        for r in p.runs:
            r.text = ""

    # Estilo base del documento
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    # Título
    t = add_paragraph(doc, "ACTA DE REUNIÓN", bold=True, size=12,
                      align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # Tabla de encabezado
    tabla = doc.add_table(rows=5, cols=2)
    tabla.style = "Normal Table"
    campos = [
        ("Empresa:", empresa),
        ("Fecha:", acta.get("fecha") or "No determinada"),
        ("Hora inicio / fin:", f"{acta.get('hora_inicio') or '—'} / {acta.get('hora_fin') or '—'}"),
        ("Lugar / Modalidad:", acta.get("modalidad") or "No determinada"),
        ("Elaborada por:", elaborado_por),
    ]
    for i, (k, v) in enumerate(campos):
        row = tabla.rows[i]
        # Altura de fila 0.5 cm
        trHeight = OxmlElement("w:trHeight")
        trHeight.set(qn("w:val"), "284")  # 0.5 cm en twips (1 cm = 567 twips)
        trHeight.set(qn("w:hRule"), "exact")
        row._tr.get_or_add_trPr().append(trHeight)
        c0, c1 = row.cells
        c0.width = Cm(4)
        # Centrado vertical en ambas celdas
        for cell in [c0, c1]:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            vAlign = OxmlElement("w:vAlign")
            vAlign.set(qn("w:val"), "center")
            tcPr.append(vAlign)
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r0 = p0.add_run(k)
        set_font(r0, size=10, bold=True)
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1 = p1.add_run(v)
        set_font(r1, size=10)

    add_paragraph(doc, space_before=6)

    # 1. Participantes
    add_section_title(doc, "1. PARTICIPANTES")
    for p in participantes:
        nombre = p.get("nombre") or p.get("etiqueta", "")
        cargo  = f" — {p['cargo']}"  if p.get("cargo")   else ""
        emp    = f" | {p['empresa']}" if p.get("empresa") else ""
        add_bullet(doc, f"{nombre}{cargo}{emp}")

    add_paragraph(doc)

    # 2. Objetivo
    add_section_title(doc, "2. OBJETIVO DE LA REUNIÓN")
    add_paragraph(doc, acta.get("objetivo", ""))
    add_paragraph(doc)

    # 3. Temas
    add_section_title(doc, "3. TEMAS TRATADOS")
    for t in acta.get("temas", []):
        p = add_bullet(doc, t.get("titulo", ""))
        # resumen con sangría
        pr = doc.add_paragraph()
        pr.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pr.paragraph_format.left_indent = Cm(0.8)
        pr.paragraph_format.line_spacing = Pt(12)
        pr.paragraph_format.space_before = Pt(0)
        pr.paragraph_format.space_after  = Pt(2)
        run = pr.add_run(t.get("resumen", ""))
        set_font(run, size=10)
    add_paragraph(doc)

    # 4. Acuerdos
    add_section_title(doc, "4. ACUERDOS Y DECISIONES")
    for a in acta.get("acuerdos", []):
        line = a.get("descripcion", "")
        if a.get("responsable"):  line += f" | Responsable: {a['responsable']}"
        if a.get("fecha_limite"): line += f" | Fecha límite: {a['fecha_limite']}"
        add_bullet(doc, line)
    add_paragraph(doc)

    # 5. Pendientes
    add_section_title(doc, "5. PUNTOS PENDIENTES / ACCIONES SIGUIENTES")
    for p in acta.get("pendientes", []):
        line = p.get("descripcion", "")
        if p.get("responsable"): line += f" | Responsable: {p['responsable']}"
        add_bullet(doc, line)
    add_paragraph(doc)

    # 6. Próxima reunión
    add_section_title(doc, "6. PRÓXIMA REUNIÓN")
    add_paragraph(doc, acta.get("proxima_reunion") or "No definida")
    add_paragraph(doc)

    # 7. Notas adicionales
    add_section_title(doc, "7. NOTAS ADICIONALES")
    for n in acta.get("notas", []):
        add_bullet(doc, n)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ── Llamadas a la API ──────────────────────────────────────────────────────────
def identificar_participantes(transcripcion, empresa):
    prompt = f"""Eres un asistente ejecutivo especializado en documentación de reuniones de negocios.
Lee esta transcripción e identifica cada participante distinto. Asigna etiqueta provisional (Participante A, B…) si no se conoce el nombre. Si alguien se presenta o es nombrado, usa el nombre real.
Responde SOLO en JSON sin texto extra ni markdown:
{{"participantes":[{{"etiqueta":"Participante A","nombre":"nombre real o null","empresa":"empresa o null","cargo":"cargo o null"}}]}}
Empresa que elabora el acta: {empresa}
Transcripción:
{transcripcion[:100000]}"""
    msg = client.messages.create(
        model="claude-sonnet-4-6", max_tokens=8000,
        messages=[{"role": "user", "content": prompt}]
    )
    text = msg.content[0].text.replace("```json", "").replace("```", "").strip()
    return json.loads(text)["participantes"]

def generar_acta(transcripcion, empresa, elaborado_por, participantes):
    lista_p = "\n".join(
        f"- {p.get('nombre') or p.get('etiqueta')}"
        + (f", {p['cargo']}" if p.get("cargo") else "")
        + (f" ({p['empresa']})" if p.get("empresa") else "")
        for p in participantes
    )
    prompt = f"""Eres un asistente ejecutivo especializado en documentación de reuniones de negocios.
Genera un acta ejecutiva formal. Responde SOLO en JSON sin markdown ni texto extra.
Formato exacto:
{{"fecha":"string o null","hora_inicio":"string o null","hora_fin":"string o null","modalidad":"string","objetivo":"string","temas":[{{"titulo":"string","resumen":"string"}}],"acuerdos":[{{"descripcion":"string","responsable":"string o null","fecha_limite":"string o null"}}],"pendientes":[{{"descripcion":"string","responsable":"string o null"}}],"proxima_reunion":"string","notas":["string"]}}
Reglas: tono formal neutro impersonal tercera persona. No uses la transcripción literalmente. No omitas acuerdos, fechas, cifras, compromisos. Ambigüedades: [inaudible] o [poco claro]. Decisiones provisionales: refléjalas.
Empresa: {empresa}
Elaborado por: {elaborado_por}
Participantes:
{lista_p}
Transcripción:
{transcripcion[:100000]}"""
    msg = client.messages.create(
        model="claude-sonnet-4-6", max_tokens=8000,
        messages=[{"role": "user", "content": prompt}]
    )
    text = msg.content[0].text.replace("```json", "").replace("```", "").strip()
    return json.loads(text)

# ── UI ─────────────────────────────────────────────────────────────────────────
st.title("📋 Generador de Actas de Reunión")

# Estado de sesión
for k, v in {"step": 1, "participantes": [], "acta": None, "transcripcion": ""}.items():
    if k not in st.session_state:
        st.session_state[k] = v

# PASO 1: Carga y empresa
if st.session_state.step >= 1:
    st.subheader("1. Transcripción y empresa")
    archivo = st.file_uploader("Carga el archivo .txt con la transcripción", type=["txt"])
    empresa = st.selectbox("Empresa", EMPRESAS)
    if archivo and st.button("Identificar participantes"):
        st.session_state.transcripcion = archivo.read().decode("utf-8")
        st.session_state.empresa = empresa
        with st.spinner("Identificando participantes…"):
            st.session_state.participantes = identificar_participantes(
                st.session_state.transcripcion, empresa)
        st.session_state.step = 2
        st.rerun()

# PASO 2: Confirmar participantes
if st.session_state.step >= 2:
    st.subheader("2. Confirma los participantes")
    participantes_edit = []
    for i, p in enumerate(st.session_state.participantes):
        with st.expander(f"{p.get('nombre') or p.get('etiqueta')}"):
            col1, col2, col3 = st.columns(3)
            nombre  = col1.text_input("Nombre",  value=p.get("nombre") or "", key=f"nom_{i}")
            cargo   = col2.text_input("Cargo",   value=p.get("cargo")  or "", key=f"car_{i}")
            emp_p   = col3.text_input("Empresa", value=p.get("empresa") or "", key=f"emp_{i}")
            participantes_edit.append({
                "etiqueta": p.get("etiqueta"), "nombre": nombre or None,
                "cargo": cargo or None, "empresa": emp_p or None
            })
    elaborado_por = st.text_input("Elaborado por (nombre completo)")
    if st.button("Generar acta") and elaborado_por:
        st.session_state.participantes = participantes_edit
        st.session_state.elaborado_por = elaborado_por
        with st.spinner("Generando acta…"):
            st.session_state.acta = generar_acta(
                st.session_state.transcripcion, st.session_state.empresa,
                elaborado_por, participantes_edit)
        st.session_state.step = 3
        st.rerun()

# PASO 3: Vista previa y descarga
if st.session_state.step >= 3 and st.session_state.acta:
    st.subheader("3. Vista previa del acta")
    acta = st.session_state.acta
    st.markdown(f"**Empresa:** {st.session_state.empresa}")
    st.markdown(f"**Fecha:** {acta.get('fecha') or 'No determinada'}")
    st.markdown(f"**Hora:** {acta.get('hora_inicio') or '—'} / {acta.get('hora_fin') or '—'}")
    st.markdown(f"**Modalidad:** {acta.get('modalidad') or 'No determinada'}")
    st.markdown(f"**Elaborada por:** {st.session_state.elaborado_por}")
    st.markdown("---")
    st.markdown("**1. PARTICIPANTES**")
    for p in st.session_state.participantes:
        n = p.get("nombre") or p.get("etiqueta")
        st.markdown(f"- {n}{' — '+p['cargo'] if p.get('cargo') else ''}{' | '+p['empresa'] if p.get('empresa') else ''}")
    st.markdown(f"**2. OBJETIVO**\n\n{acta.get('objetivo','')}")
    st.markdown("**3. TEMAS TRATADOS**")
    for t in acta.get("temas", []):
        st.markdown(f"- **{t['titulo']}**\n\n  {t['resumen']}")
    st.markdown("**4. ACUERDOS Y DECISIONES**")
    for a in acta.get("acuerdos", []):
        line = a["descripcion"]
        if a.get("responsable"):  line += f" | Resp.: {a['responsable']}"
        if a.get("fecha_limite"): line += f" | Fecha: {a['fecha_limite']}"
        st.markdown(f"- {line}")
    st.markdown("**5. PUNTOS PENDIENTES**")
    for p in acta.get("pendientes", []):
        line = p["descripcion"]
        if p.get("responsable"): line += f" | Resp.: {p['responsable']}"
        st.markdown(f"- {line}")
    st.markdown(f"**6. PRÓXIMA REUNIÓN**\n\n{acta.get('proxima_reunion') or 'No definida'}")
    st.markdown("**7. NOTAS ADICIONALES**")
    for n in acta.get("notas", []):
        st.markdown(f"- {n}")
    st.markdown("---")

    nombre_archivo = st.text_input("Nombre del archivo (sin extensión)", value="Acta_Reunion")
    docx_buf = build_docx(acta, st.session_state.participantes,
                          st.session_state.empresa, st.session_state.elaborado_por)
    st.download_button(
        label="⬇ Descargar .docx",
        data=docx_buf,
        file_name=f"{nombre_archivo}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    if st.button("🔄 Nueva acta"):
        for k in ["step", "participantes", "acta", "transcripcion"]:
            del st.session_state[k]
        st.rerun()
